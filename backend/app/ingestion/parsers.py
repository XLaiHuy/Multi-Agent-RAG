"""
Multi-Format Canonical Document Parsers.
Converts PDF, Scanned PDF (with OCR gating), Markdown, JSON, DOCX, and Images
into unified CanonicalDocument representations with pages, blocks, bounding boxes, and section paths.
"""
import io
import json
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

import fitz  # PyMuPDF
from backend.app.domain.canonical import (
    CanonicalDocument, CanonicalPage, CanonicalBlock, BlockType, BoundingBox
)
from backend.app.providers.interfaces import OCRProvider
from backend.app.core.config import get_settings

logger = logging.getLogger("document_parsers")


class DoclingPDFParserAdapter:
    """Optional Docling parser adapter for advanced layout, tables, and OCR parsing."""

    @staticmethod
    def is_available() -> bool:
        try:
            import docling  # noqa: F401
            return True
        except ImportError:
            return False

    @staticmethod
    def parse_pdf(file_path: Path, doc_id: str) -> Optional[CanonicalDocument]:
        """Parses PDF with Docling if available; returns None on failure or missing dependency."""
        if not DoclingPDFParserAdapter.is_available():
            return None
        try:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            res = converter.convert(str(file_path))
            doc = res.document
            page_blocks_map: Dict[int, List[CanonicalBlock]] = {}

            for item, _ in doc.iterate_items():
                text = getattr(item, "text", "") or ""
                text = text.strip()
                if not text:
                    continue
                p_num = 1
                if hasattr(item, "prov") and item.prov and len(item.prov) > 0:
                    p_num = getattr(item.prov[0], "page_no", 1) or 1

                b_type = BlockType.TABLE if "table" in str(type(item)).lower() else BlockType.PARAGRAPH
                b_id = f"{doc_id}_p{p_num}_docling_{len(page_blocks_map.get(p_num, []))}"
                block = CanonicalBlock(
                    block_id=b_id,
                    block_type=b_type,
                    text=text,
                    page_number=p_num,
                    metadata={"parser": "docling"},
                )
                page_blocks_map.setdefault(p_num, []).append(block)

            canonical_pages: List[CanonicalPage] = []
            for p_num in sorted(page_blocks_map.keys()):
                canonical_pages.append(CanonicalPage(page_number=p_num, blocks=page_blocks_map[p_num]))

            if canonical_pages:
                return CanonicalDocument(
                    doc_id=doc_id,
                    title=file_path.name,
                    doc_type="pdf",
                    pages=canonical_pages,
                    metadata={"parsed_by": "docling"},
                )
        except Exception as e:
            logger.warning(f"[DoclingAdapter] Docling parsing fallback to NativePDFParser: {e}")
        return None



class OCRGatingAnalyzer:
    """
    Analyzes PDF pages using multi-signal heuristics to determine if OCR is required:
    - Text character density per unit area
    - Non-printable / malformed character ratio
    - Image coverage ratio on the page
    - Presence of selectable text objects
    """

    @staticmethod
    def should_ocr_page(page: fitz.Page, min_char_threshold: int = 40) -> Tuple[bool, str]:
        rect = page.rect
        page_area = rect.width * rect.height
        if page_area <= 0:
            return True, "invalid_page_dimensions"

        text = page.get_text("text").strip()
        char_count = len(text)

        # 1. Very low text count
        if char_count < min_char_threshold:
            # Check if page contains images
            image_list = page.get_images()
            if image_list:
                return True, f"low_text_with_images (chars={char_count}, imgs={len(image_list)})"
            return True, f"sparse_or_empty_page (chars={char_count})"

        # 2. Check for malformed / non-printable unicode character ratio
        non_printable = sum(1 for c in text if not c.isprintable() and c not in "\n\r\t")
        if char_count > 0 and (non_printable / char_count) > 0.15:
            return True, f"high_malformed_char_ratio ({non_printable}/{char_count})"

        # 3. Check image coverage vs text density
        images = page.get_images()
        if len(images) >= 1 and char_count < 150:
            return True, f"heavy_image_low_text (imgs={len(images)}, chars={char_count})"

        return False, "native_text_acceptable"


class NativePDFParser:
    """Extracts structured blocks and layout coordinates from vector/text PDF pages."""

    @staticmethod
    def parse_pdf(file_path: Path, doc_id: str) -> CanonicalDocument:
        doc = fitz.open(str(file_path))
        canonical_pages: List[CanonicalPage] = []
        current_section_path: List[str] = []

        for page_idx in range(len(doc)):
            page = doc[page_idx]
            page_num = page_idx + 1
            page_blocks: List[CanonicalBlock] = []

            # Extract layout blocks with bounding boxes: (x0, y0, x1, y1, text, block_no, block_type)
            # block_type == 0: Text block, block_type == 1: Image block
            raw_blocks = page.get_text("blocks")

            for b_idx, b in enumerate(raw_blocks):
                x0, y0, x1, y1, text, block_no, b_type = b[:7]
                text = text.strip()
                if not text:
                    continue

                bbox = BoundingBox(x0=round(x0, 2), y0=round(y0, 2), x1=round(x1, 2), y1=round(y1, 2))
                
                # Heading detection heuristic (short text, starts with Article/Section/Chapter/digits)
                is_heading = False
                if len(text.split("\n")) <= 2 and len(text) < 120:
                    if re.match(r"(?i)^(Article|Section|Clause|Chapter|Part|\d+\.|\d+\.\d+)", text):
                        is_heading = True
                        current_section_path = [text]

                block_type = BlockType.HEADING if is_heading else BlockType.PARAGRAPH
                block_id = f"{doc_id}_p{page_num}_b{b_idx}"

                canonical_blocks = CanonicalBlock(
                    block_id=block_id,
                    block_type=block_type,
                    text=text,
                    page_number=page_num,
                    section_path=list(current_section_path),
                    bbox=bbox,
                    source_offset_start=0,
                    source_offset_end=len(text),
                )
                page_blocks.append(canonical_blocks)

            canonical_pages.append(
                CanonicalPage(
                    page_number=page_num,
                    width=round(page.rect.width, 2),
                    height=round(page.rect.height, 2),
                    blocks=page_blocks,
                    is_scanned=False,
                )
            )

        doc.close()
        return CanonicalDocument(
            doc_id=doc_id,
            title=file_path.name,
            doc_type="pdf",
            pages=canonical_pages,
            metadata={"filename": file_path.name, "page_count": len(canonical_pages)},
        )


class MarkdownParser:
    """Parses Markdown files preserving heading hierarchy into section paths and tables."""

    @staticmethod
    def parse_markdown(file_path: Path, doc_id: str) -> CanonicalDocument:
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        lines = text.split("\n")
        
        blocks: List[CanonicalBlock] = []
        current_section_stack: List[Tuple[int, str]] = [] # [(level, text)]
        current_paragraph: List[str] = []
        block_counter = 0

        def flush_paragraph():
            nonlocal block_counter, current_paragraph
            if current_paragraph:
                p_text = "\n".join(current_paragraph).strip()
                if p_text:
                    section_path = [s[1] for s in current_section_stack]
                    b_id = f"{doc_id}_b{block_counter}"
                    blocks.append(
                        CanonicalBlock(
                            block_id=b_id,
                            block_type=BlockType.PARAGRAPH,
                            text=p_text,
                            page_number=1,
                            section_path=section_path,
                        )
                    )
                    block_counter += 1
                current_paragraph = []

        in_code_block = False
        code_lines = []

        for line in lines:
            trimmed = line.strip()
            if trimmed.startswith("```"):
                if in_code_block:
                    # Closing code block
                    code_lines.append(line)
                    c_text = "\n".join(code_lines)
                    b_id = f"{doc_id}_b{block_counter}"
                    blocks.append(
                        CanonicalBlock(
                            block_id=b_id,
                            block_type=BlockType.CODE,
                            text=c_text,
                            page_number=1,
                            section_path=[s[1] for s in current_section_stack],
                        )
                    )
                    block_counter += 1
                    code_lines = []
                    in_code_block = False
                else:
                    flush_paragraph()
                    in_code_block = True
                    code_lines.append(line)
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            # Heading check
            heading_match = re.match(r"^(#{1,6})\s+(.*)", trimmed)
            if heading_match:
                flush_paragraph()
                level = len(heading_match.group(1))
                heading_title = heading_match.group(2).strip()

                # Pop headings of equal or deeper level
                while current_section_stack and current_section_stack[-1][0] >= level:
                    current_section_stack.pop()
                current_section_stack.append((level, heading_title))

                b_id = f"{doc_id}_b{block_counter}"
                blocks.append(
                    CanonicalBlock(
                        block_id=b_id,
                        block_type=BlockType.HEADING,
                        text=heading_title,
                        page_number=1,
                        section_path=[s[1] for s in current_section_stack],
                    )
                )
                block_counter += 1
            elif not trimmed:
                flush_paragraph()
            else:
                current_paragraph.append(line)

        flush_paragraph()

        page = CanonicalPage(page_number=1, blocks=blocks)
        return CanonicalDocument(
            doc_id=doc_id,
            title=file_path.name,
            doc_type="markdown",
            pages=[page],
            metadata={"filename": file_path.name, "block_count": len(blocks)},
        )


class JSONParser:
    """
    Parses structured JSON contracts into canonical hierarchical blocks.
    Preserves exact JSON paths (e.g. 'contract.termination.notice_period').
    """

    @staticmethod
    def parse_json(file_path: Path, doc_id: str) -> CanonicalDocument:
        content = file_path.read_text(encoding="utf-8", errors="ignore")
        data = json.loads(content)
        blocks: List[CanonicalBlock] = []
        block_counter = 0

        def traverse(node: Any, current_path: List[str]):
            nonlocal block_counter
            if isinstance(node, dict):
                for k, v in node.items():
                    traverse(v, current_path + [str(k)])
            elif isinstance(node, list):
                for idx, item in enumerate(node):
                    traverse(item, current_path + [f"[{idx}]"])
            else:
                # Leaf node: convert to text with JSON Path
                json_path_str = ".".join(current_path)
                val_str = str(node)
                text = f"**{json_path_str}**: {val_str}"
                
                b_id = f"{doc_id}_b{block_counter}"
                blocks.append(
                    CanonicalBlock(
                        block_id=b_id,
                        block_type=BlockType.CLAUSE,
                        text=text,
                        page_number=1,
                        section_path=current_path,
                        metadata={"json_path": json_path_str, "value": val_str},
                    )
                )
                block_counter += 1

        traverse(data, ["root"])
        page = CanonicalPage(page_number=1, blocks=blocks)
        return CanonicalDocument(
            doc_id=doc_id,
            title=file_path.name,
            doc_type="json",
            pages=[page],
            metadata={"filename": file_path.name, "json_keys_count": len(blocks)},
        )


class DocxParser:
    """Parses Microsoft Word (.docx) documents preserving headings, paragraphs, and tables."""

    @staticmethod
    def parse_docx(file_path: Path, doc_id: str) -> CanonicalDocument:
        import docx
        doc = docx.Document(file_path)
        blocks: List[CanonicalBlock] = []
        current_section_path: List[str] = []
        block_counter = 0

        for p in doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue

            if p.style.name.startswith("Heading"):
                current_section_path = [txt]
                b_type = BlockType.HEADING
            else:
                b_type = BlockType.PARAGRAPH

            b_id = f"{doc_id}_b{block_counter}"
            blocks.append(
                CanonicalBlock(
                    block_id=b_id,
                    block_type=b_type,
                    text=txt,
                    page_number=1,
                    section_path=list(current_section_path),
                )
            )
            block_counter += 1

        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                table_rows.append(row_cells)
            
            if table_rows:
                # Convert table rows to markdown representation
                md_table_lines = ["| " + " | ".join(table_rows[0]) + " |"]
                md_table_lines.append("| " + " | ".join(["---"] * len(table_rows[0])) + " |")
                for r in table_rows[1:]:
                    md_table_lines.append("| " + " | ".join(r) + " |")
                
                table_text = "\n".join(md_table_lines)
                b_id = f"{doc_id}_b{block_counter}"
                blocks.append(
                    CanonicalBlock(
                        block_id=b_id,
                        block_type=BlockType.TABLE,
                        text=table_text,
                        page_number=1,
                        section_path=list(current_section_path),
                        table_data=table_rows,
                    )
                )
                block_counter += 1

        page = CanonicalPage(page_number=1, blocks=blocks)
        return CanonicalDocument(
            doc_id=doc_id,
            title=file_path.name,
            doc_type="docx",
            pages=[page],
            metadata={"filename": file_path.name},
        )


class MasterDocumentParser:
    """Unified Document Parser dispatching to appropriate format parser."""

    @staticmethod
    def parse(file_path: Path, doc_id: str, ocr_provider: Optional[OCRProvider] = None) -> CanonicalDocument:
        ext = file_path.suffix.lower()
        if ext == ".pdf":
            # 1. If Docling is available and configured, attempt Docling layout parsing
            settings = get_settings()
            if getattr(settings, "use_docling_parser", False) and DoclingPDFParserAdapter.is_available():
                docling_doc = DoclingPDFParserAdapter.parse_pdf(file_path, doc_id)
                if docling_doc:
                    return docling_doc

            # 2. Native PyMuPDF extraction
            native_doc = NativePDFParser.parse_pdf(file_path, doc_id)
            # If native extraction yielded ample text, return native
            total_chars = sum(len(b.text) for b in native_doc.get_all_blocks())
            if total_chars >= 100:
                return native_doc
            
            # Scanned / empty PDF fallback with OCR if provider available
            if ocr_provider:
                fitz_doc = fitz.open(str(file_path))
                ocr_pages = []
                for p_idx in range(len(fitz_doc)):
                    page = fitz_doc[p_idx]
                    should_ocr, _ = OCRGatingAnalyzer.should_ocr_page(page)
                    if should_ocr:
                        pix = page.get_pixmap(dpi=200)
                        img_bytes = pix.tobytes(output="png")
                        extracted_text = ocr_provider.extract_from_image_bytes(img_bytes, mime_type="image/png")
                        b_id = f"{doc_id}_p{p_idx+1}_ocr_b0"
                        block = CanonicalBlock(
                            block_id=b_id,
                            block_type=BlockType.PARAGRAPH,
                            text=extracted_text,
                            page_number=p_idx + 1,
                            metadata={"ocr_processed": True},
                        )
                        ocr_pages.append(CanonicalPage(page_number=p_idx + 1, blocks=[block], is_scanned=True))
                    else:
                        # Use native page
                        ocr_pages.append(native_doc.pages[p_idx])
                fitz_doc.close()
                return CanonicalDocument(
                    doc_id=doc_id, title=file_path.name, doc_type="pdf", pages=ocr_pages, metadata={"is_scanned": True}
                )
            return native_doc

        elif ext in [".md", ".txt"]:
            return MarkdownParser.parse_markdown(file_path, doc_id)
        elif ext == ".json":
            return JSONParser.parse_json(file_path, doc_id)
        elif ext in [".docx", ".doc"]:
            return DocxParser.parse_docx(file_path, doc_id)
        elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp"]:
            if ocr_provider:
                img_bytes = file_path.read_bytes()
                mime_type = "image/jpeg" if ext in [".jpg", ".jpeg"] else "image/png"
                text = ocr_provider.extract_from_image_bytes(img_bytes, mime_type=mime_type)
                block = CanonicalBlock(
                    block_id=f"{doc_id}_img_b0",
                    block_type=BlockType.PARAGRAPH,
                    text=text,
                    page_number=1,
                    metadata={"image_ocr": True},
                )
                return CanonicalDocument(
                    doc_id=doc_id,
                    title=file_path.name,
                    doc_type="image",
                    pages=[CanonicalPage(page_number=1, blocks=[block], is_scanned=True)],
                )
            else:
                raise ValueError(f"OCR provider required for image file {file_path.name}")
        else:
            raise ValueError(f"Unsupported file format: {ext}")
