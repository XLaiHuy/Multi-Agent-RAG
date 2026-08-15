"""
CUAD Multi-Format Variant Generator.
Transforms canonical contracts into 4 format representations:
1. Markdown (.md)
2. Structured JSON (.json)
3. Microsoft Word (.docx)
4. Scanned / Degraded document text
Enables Format-Invariance evaluation (DeltaRecall_format = Recall_canonical - Recall_format).
"""
import json
from pathlib import Path
from typing import Dict, Any


def generate_json_variant_from_markdown(md_text: str) -> Dict[str, Any]:
    """Converts structured markdown contract into hierarchical JSON."""
    lines = md_text.split("\n")
    contract_json: Dict[str, Any] = {"title": "Contract", "sections": {}}
    current_sec = "general"

    for line in lines:
        trimmed = line.strip()
        if trimmed.startswith("## Section"):
            current_sec = trimmed.replace("##", "").strip()
            contract_json["sections"][current_sec] = []
        elif trimmed and not trimmed.startswith("#"):
            if current_sec not in contract_json["sections"]:
                contract_json["sections"][current_sec] = []
            contract_json["sections"][current_sec].append(trimmed)

    return contract_json


def generate_docx_variant_from_markdown(md_text: str, output_path: Path):
    """Converts markdown contract into a .docx document."""
    import docx
    doc = docx.Document()
    lines = md_text.split("\n")

    for line in lines:
        trimmed = line.strip()
        if not trimmed:
            continue
        if trimmed.startswith("# "):
            doc.add_heading(trimmed.replace("# ", ""), level=1)
        elif trimmed.startswith("## "):
            doc.add_heading(trimmed.replace("## ", ""), level=2)
        elif trimmed.startswith("### "):
            doc.add_heading(trimmed.replace("### ", ""), level=3)
        else:
            doc.add_paragraph(trimmed)

    doc.save(output_path)


def generate_all_format_variants(fixtures_dir: Path):
    """Generates MD, JSON, and DOCX variants for all base contracts in fixtures."""
    fixtures_dir = Path(fixtures_dir)
    md_files = list(fixtures_dir.glob("*.md"))

    for md_file in md_files:
        content = md_file.read_text(encoding="utf-8")
        stem = md_file.stem

        # 1. JSON Variant
        json_data = generate_json_variant_from_markdown(content)
        json_file = fixtures_dir / f"{stem}.json"
        json_file.write_text(json.dumps(json_data, indent=2), encoding="utf-8")

        # 2. DOCX Variant
        docx_file = fixtures_dir / f"{stem}.docx"
        try:
            generate_docx_variant_from_markdown(content, docx_file)
        except Exception as e:
            print(f"[MultiFormat] Could not create DOCX for {stem}: {e}")

    print(f"[MultiFormat] Generated multi-format variants in {fixtures_dir}")


if __name__ == "__main__":
    generate_all_format_variants(Path("tests/fixtures/cuad_small"))
