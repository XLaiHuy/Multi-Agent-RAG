"""
Multi-Format Robustness Benchmark Suite.
Generates and parses official CUAD contract formats:
- Canonical Text (.txt)
- Markdown (.md)
- Structured JSON (.json)
- Microsoft Word (.docx)
- Vector / Text PDF (.pdf)
Measures: Parse Success Rate, Character Preservation, Structure/Block Extraction,
and Downstream Retrieval (Recall@5, MRR, nDCG@5).
Saves raw artifacts to evaluation/runs/<run_id>/multiformat/
"""
import io
import json
import time
import uuid
import datetime
from pathlib import Path
from typing import Dict, Any, List

import re
import numpy as np
import docx  # python-docx
import fitz  # PyMuPDF

from backend.app.ingestion.parsers import MasterDocumentParser
from backend.app.ingestion.chunker import StructureAwareParentChildChunker
from backend.app.retrieval.bm25 import BM25Retriever
from evaluation.metrics.retrieval_metrics import compute_recall_at_k, compute_reciprocal_rank, compute_ndcg_at_k

RUNS_DIR = Path("evaluation/runs")
REPORTS_DIR = Path("evaluation/reports")
MANIFEST_PATH = Path("evaluation/manifests/cuad_official_manifest.json")
CONTRACTS_DIR = Path("evaluation/datasets/cuad/processed/contracts")
FORMATS_DIR = Path("evaluation/datasets/cuad/processed/format_variants")

RUNS_DIR.mkdir(parents=True, exist_ok=True)
REPORTS_DIR.mkdir(parents=True, exist_ok=True)
FORMATS_DIR.mkdir(parents=True, exist_ok=True)


def generate_format_variants(source_contract_id: str, raw_text: str, title: str) -> Dict[str, Path]:
    variants = {}

    # 1. Canonical Text
    txt_path = FORMATS_DIR / f"{source_contract_id}.txt"
    txt_path.write_text(raw_text, encoding="utf-8")
    variants["TXT"] = txt_path

    # 2. Markdown
    md_path = FORMATS_DIR / f"{source_contract_id}.md"
    md_path.write_text(f"# {title}\n\n{raw_text}", encoding="utf-8")
    variants["MD"] = md_path

    # 3. Structured JSON
    json_path = FORMATS_DIR / f"{source_contract_id}.json"
    paragraphs = raw_text.split("\n\n")
    json_data = {
        "title": title,
        "contract_id": source_contract_id,
        "sections": [
            {"section_id": f"s_{i+1}", "title": f"Clause {i+1}", "content": p.strip()}
            for i, p in enumerate(paragraphs) if p.strip()
        ]
    }
    json_path.write_text(json.dumps(json_data, indent=2), encoding="utf-8")
    variants["JSON"] = json_path

    # 4. DOCX
    docx_path = FORMATS_DIR / f"{source_contract_id}.docx"
    doc = docx.Document()
    doc.add_heading(title, level=0)
    for p in paragraphs:
        if p.strip():
            doc.add_paragraph(p.strip())
    doc.save(docx_path)
    variants["DOCX"] = docx_path

    # 5. PDF
    pdf_path = FORMATS_DIR / f"{source_contract_id}.pdf"
    pdf_doc = fitz.open()
    # Paginate into A4 pages
    lines = raw_text.split("\n")
    cur_lines = []
    for line in lines:
        cur_lines.append(line)
        if len(cur_lines) >= 45:
            p = pdf_doc.new_page(width=595, height=842)
            p.insert_textbox(fitz.Rect(40, 40, 555, 800), "\n".join(cur_lines), fontsize=9.5)
            cur_lines = []
    if cur_lines:
        p = pdf_doc.new_page(width=595, height=842)
        p.insert_textbox(fitz.Rect(40, 40, 555, 800), "\n".join(cur_lines), fontsize=9.5)
    pdf_doc.save(str(pdf_path))
    variants["PDF"] = pdf_path

    return variants


def run_multiformat_benchmark() -> Dict[str, Any]:
    run_id = f"multiformat_run_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}"
    run_dir = RUNS_DIR / run_id / "multiformat"
    run_dir.mkdir(parents=True, exist_ok=True)

    print(f"[Multi-Format Benchmark] Initializing Run ID: {run_id}...")

    manifest_data = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    contracts = manifest_data["contracts"]
    queries = manifest_data["queries"]

    target_contract = contracts[0]
    target_id = target_contract["source_contract_id"]
    raw_text = (CONTRACTS_DIR / target_contract["filename"]).read_text(encoding="utf-8")
    title = target_contract["original_title"]

    print(f"[Multi-Format Benchmark] Generating variants for: {target_id}...")
    variants = generate_format_variants(target_id, raw_text, title)

    chunker = StructureAwareParentChildChunker(child_target_tokens=200, parent_target_tokens=800)
    target_queries = [q for q in queries if q["source_contract_id"] == target_id][:5]

    format_results = {}
    raw_rows = []

    for fmt, path in variants.items():
        t0 = time.perf_counter()
        try:
            canonical_doc = MasterDocumentParser.parse(path, doc_id=target_id)
            c_chunks, p_chunks = chunker.chunk_canonical_document(canonical_doc, doc_version=1)
            
            # Build search index
            bm25 = BM25Retriever()
            bm25.build_index([c.chunk_id for c in c_chunks], [c.text for c in c_chunks], [c.metadata for c in c_chunks])

            recalls, mrrs, ndcgs = [], [], []
            for q in target_queries:
                hits = bm25.search(q["question"], top_k=5)
                gold_words = set(re.findall(r"\b\w+\b", q["gold_evidence"].lower()[:80]))
                hit_match = False
                rank_hit = 0
                for rank, (hid, score, hmeta) in enumerate(hits, 1):
                    c_text = hmeta.get("parent_text", "") + " " + hmeta.get("text", "")
                    chunk_words = set(re.findall(r"\b\w+\b", c_text.lower()))
                    if len(gold_words.intersection(chunk_words)) >= min(len(gold_words), 3):
                        hit_match = True
                        rank_hit = rank
                        break
                recalls.append(1.0 if hit_match else 0.0)
                mrrs.append(1.0 / rank_hit if hit_match else 0.0)
                ndcgs.append(1.0 / np.log2(rank_hit + 1) if hit_match else 0.0)

            elapsed_ms = (time.perf_counter() - t0) * 1000
            total_chars = sum(len(b.text) for p in canonical_doc.pages for b in p.blocks)
            total_blocks = sum(len(p.blocks) for p in canonical_doc.pages)

            r_score = sum(recalls) / len(recalls) if recalls else 0.0
            mrr_score = sum(mrrs) / len(mrrs) if mrrs else 0.0
            ndcg_score = sum(ndcgs) / len(ndcgs) if ndcgs else 0.0

            res = {
                "format": fmt,
                "file_path": str(path),
                "parse_success": True,
                "pages_extracted": len(canonical_doc.pages),
                "blocks_extracted": total_blocks,
                "child_chunks_indexed": len(c_chunks),
                "parent_chunks_indexed": len(p_chunks),
                "chars_extracted": total_chars,
                "Recall@5": round(r_score, 4),
                "MRR": round(mrr_score, 4),
                "nDCG@5": round(ndcg_score, 4),
                "latency_ms": round(elapsed_ms, 2)
            }
            format_results[fmt] = res
            raw_rows.append(res)
            print(f"  Format [{fmt:<4}]: Pages={len(canonical_doc.pages):2d} | Blocks={total_blocks:3d} | R@5={r_score:.3f} | Latency={elapsed_ms:.1f}ms")

        except Exception as e:
            elapsed_ms = (time.perf_counter() - t0) * 1000
            err_res = {
                "format": fmt,
                "file_path": str(path),
                "parse_success": False,
                "error": str(e),
                "latency_ms": round(elapsed_ms, 2)
            }
            format_results[fmt] = err_res
            raw_rows.append(err_res)
            print(f"  Format [{fmt:<4}]: FAILED - {e}")

    # Save raw outputs
    with open(run_dir / "format_metrics.jsonl", "w", encoding="utf-8") as f:
        for r in raw_rows:
            f.write(json.dumps(r) + "\n")

    summary = {
        "run_id": run_id,
        "benchmark_timestamp": time.strftime("%Y-%m-%dT%H:%M:%SZ"),
        "contract_id": target_id,
        "queries_count": len(target_queries),
        "results": format_results
    }

    with open(run_dir / "summary.json", "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2)

    with open(REPORTS_DIR / "multiformat_benchmark_results.json", "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2)

    print(f"\n[Multi-Format Benchmark] Complete! Summary saved to {REPORTS_DIR / 'multiformat_benchmark_results.json'}")
    return summary


if __name__ == "__main__":
    run_multiformat_benchmark()
