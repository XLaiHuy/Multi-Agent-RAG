"""
Document Loader Đa Định DạngSiêu Tốc (Fast Multi-Format Loader):
- PDF (.pdf): Thử PyMuPDF native C++ text extraction trước (<10ms), fallback sang Vision OCR nếu scanned
- Word (.docx, .doc): Trích xuất tiêu đề, đoạn văn và bảng biểu định dạng Markdown
- Excel / CSV (.xlsx, .xls, .csv): Đọc các sheet và định dạng thành bảng Markdown | col 1 | col 2 |
- JSON (.json): Phân tích cấu trúc dữ liệu JSON thành Markdown
- Text / Markdown (.md, .txt)
- Images (.png, .jpg, .jpeg, .webp, .bmp, .tiff) qua Gemini Multimodal Vision OCR
"""
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, List

import fitz  # PyMuPDF


@dataclass
class RawDocument:
    doc_id: str
    source: str
    text: str
    metadata: dict = field(default_factory=dict)


def load_text_file(path: Path) -> RawDocument:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    text = path.read_text(encoding="utf-8", errors="ignore")
    if not text.strip():
        raise ValueError(f"File is empty: {path}")

    return RawDocument(
        doc_id=path.stem,
        source=path.name,
        text=text,
        metadata={
            "filename": path.name,
            "source": path.name,
            "file_type": path.suffix.lower(),
            "char_count": len(text),
        },
    )


def load_json_file(path: Path) -> RawDocument:
    """Load và định dạng file JSON thành Markdown rành mạch."""
    path = Path(path)
    raw_data = path.read_text(encoding="utf-8", errors="ignore")
    try:
        parsed = json.loads(raw_data)
        formatted_md = f"# Tài liệu JSON: {path.name}\n\n```json\n" + json.dumps(parsed, ensure_ascii=False, indent=2) + "\n```"
    except Exception:
        formatted_md = raw_data

    return RawDocument(
        doc_id=path.stem,
        source=path.name,
        text=formatted_md,
        metadata={
            "filename": path.name,
            "source": path.name,
            "file_type": ".json",
            "char_count": len(formatted_md),
        },
    )


def load_word_file(path: Path) -> RawDocument:
    """Load file Word (.docx, .doc) và chuyển đổi thành Markdown cấu trúc."""
    path = Path(path)
    try:
        import docx
    except ImportError:
        raise ImportError("Cần cài đặt `python-docx` để đọc file Word (.docx). Run: pip install python-docx")

    doc = docx.Document(path)
    md_lines = []

    # Read paragraphs
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        if p.style.name.startswith("Heading 1"):
            md_lines.append(f"# {txt}\n")
        elif p.style.name.startswith("Heading 2"):
            md_lines.append(f"## {txt}\n")
        elif p.style.name.startswith("Heading 3"):
            md_lines.append(f"### {txt}\n")
        else:
            md_lines.append(f"{txt}\n")

    # Read tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cell_texts = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append(cell_texts)
        
        if rows:
            # Header
            header = rows[0]
            md_lines.append("| " + " | ".join(header) + " |")
            md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
            for r in rows[1:]:
                md_lines.append("| " + " | ".join(r) + " |")
            md_lines.append("\n")

    full_text = "\n".join(md_lines).strip()
    if not full_text:
        full_text = f"Tài liệu Word {path.name} (Rống)"

    return RawDocument(
        doc_id=path.stem,
        source=path.name,
        text=full_text,
        metadata={
            "filename": path.name,
            "source": path.name,
            "file_type": path.suffix.lower(),
            "char_count": len(full_text),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        },
    )


def load_excel_csv_file(path: Path) -> RawDocument:
    """Load file Excel (.xlsx, .xls) hoặc CSV (.csv) và định dạng thành bảng Markdown."""
    path = Path(path)
    ext = path.suffix.lower()
    md_lines = [f"# Bảng dữ liệu: {path.name}\n"]

    try:
        import pandas as pd
        if ext == ".csv":
            df = pd.read_csv(path)
            md_lines.append(df.to_markdown(index=False) or "")
        else:
            # Excel file with multiple sheets
            excel = pd.ExcelFile(path)
            for sheet in excel.sheet_names:
                df = pd.read_excel(excel, sheet_name=sheet)
                md_lines.append(f"## Sheet: {sheet}\n")
                md_lines.append(df.to_markdown(index=False) or "")
                md_lines.append("\n")
    except Exception as e:
        # Fallback reading CSV as text
        md_lines.append(path.read_text(encoding="utf-8", errors="ignore"))

    full_text = "\n".join(md_lines).strip()

    return RawDocument(
        doc_id=path.stem,
        source=path.name,
        text=full_text,
        metadata={
            "filename": path.name,
            "source": path.name,
            "file_type": ext,
            "char_count": len(full_text),
        },
    )


def load_image_document(path: Path) -> RawDocument:
    """Load file ảnh (.png, .jpg, .jpeg, .webp, .bmp) bằng Gemini Multimodal Vision OCR."""
    from app.ingestion.ocr import get_ocr_engine

    path = Path(path)
    engine = get_ocr_engine()
    print(f"  [Loader] Running Multimodal OCR for image: {path.name}...")
    markdown_text = engine.extract_from_image_file(path)

    if not markdown_text or markdown_text.startswith("[OCR Error"):
        raise ValueError(f"Không thể trích xuất nội dung từ ảnh {path.name}: {markdown_text}")

    return RawDocument(
        doc_id=path.stem,
        source=path.name,
        text=markdown_text,
        metadata={
            "filename": path.name,
            "source": path.name,
            "file_type": "image",
            "image_format": path.suffix.lower(),
            "char_count": len(markdown_text),
            "ocr_engine": "gemini_multimodal_vision",
        },
    )


def load_pdf_fast_native(path: Path) -> Optional[RawDocument]:
    """
    Thử trích xuất văn bản siêu tốc bằng PyMuPDF C++ Native (<10ms).
    Trả về None nếu phát hiện PDF là tài liệu scan/ảnh (không có text layer).
    """
    doc = fitz.open(str(path))
    total_pages = len(doc)
    page_texts = []
    scanned_pages = 0

    for i in range(total_pages):
        text = doc[i].get_text("text").strip()
        if len(text) < 40: # Trang bị nghi vấn là scanned image
            scanned_pages += 1
        page_texts.append(f"--- Trang {i+1} ---\n" + text)

    # Nếu trên 50% số trang là scanned image -> Chuyển sang Vision OCR
    if scanned_pages / max(total_pages, 1) > 0.5:
        doc.close()
        return None

    full_text = "\n\n".join(page_texts).strip()
    doc.close()

    if len(full_text) < 50:
        return None

    return RawDocument(
        doc_id=path.stem,
        source=path.name,
        text=full_text,
        metadata={
            "filename": path.name,
            "source": path.name,
            "file_type": ".pdf",
            "pdf_type": "native_text_fast",
            "page_count": total_pages,
            "char_count": len(full_text),
        },
    )


def load_pdf_file(path: Path) -> RawDocument:
    """
    Tối ưu hóa quy trình nạp PDF 3 Tầng (Fast-Track PDF Pipeline):
    1. Thử PyMuPDF Native (<10ms): Trích xuất ngay lập tức nếu là PDF có text.
    2. Thử pdf-inspector: Giữ nguyên cấu trúc Markdown nếu có bảng/code.
    3. Thử Gemini Vision OCR Engine: Xử lý PDF dạng ảnh scan / biểu đồ.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    # Tầng 1: Native Fast Extraction (<10ms)
    try:
        fast_doc = load_pdf_fast_native(path)
        if fast_doc and len(fast_doc.text) >= 100:
            print(f"  ⚡ [Fast PDF Loader] {path.name}: Native text extracted in <10ms ({fast_doc.metadata['char_count']} chars).")
            return fast_doc
    except Exception as e:
        print(f"  [Fast PDF Loader] {path.name} native bypass ({e}).")

    # Tầng 2: Thử pdf-inspector
    try:
        import pdf_inspector
        result = pdf_inspector.process_pdf(str(path))
        markdown_text = result.markdown or ""
        if len(markdown_text.strip()) >= 100:
            return RawDocument(
                doc_id=path.stem, source=path.name, text=markdown_text,
                metadata={"filename": path.name, "source": path.name, "file_type": ".pdf", "pdf_type": "pdf_inspector"}
            )
    except Exception:
        pass

    # Tầng 3: Vision OCR (dành cho scanned PDF)
    from app.ingestion.ocr import get_ocr_engine
    engine = get_ocr_engine()
    markdown_text = engine.extract_scanned_pdf(path)

    return RawDocument(
        doc_id=path.stem,
        source=path.name,
        text=markdown_text or f"Tài liệu PDF {path.name}",
        metadata={"filename": path.name, "source": path.name, "file_type": ".pdf", "pdf_type": "vision_ocr"}
    )


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"}


def load_document(path: Path) -> RawDocument:
    """Bộ định tuyến nạp tài liệu đa định dạng."""
    path = Path(path)
    ext = path.suffix.lower()

    if ext in [".md", ".txt"]:
        return load_text_file(path)
    elif ext == ".json":
        return load_json_file(path)
    elif ext in [".docx", ".doc"]:
        return load_word_file(path)
    elif ext in [".xlsx", ".xls", ".csv"]:
        return load_excel_csv_file(path)
    elif ext == ".pdf":
        return load_pdf_file(path)
    elif ext in IMAGE_EXTENSIONS:
        return load_image_document(path)
    else:
        raise ValueError(f"Không hỗ trợ định dạng '{ext}' cho file: {path.name}")
