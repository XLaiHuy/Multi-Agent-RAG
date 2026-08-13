import pytest
from pathlib import Path
from PIL import Image, ImageDraw
import json
import io

from app.ingestion.loader import load_document, load_text_file, load_image_document


def test_load_text_file(tmp_path: Path):
    sample_file = tmp_path / "sample.md"
    sample_file.write_text("# Test Title\nThis is a sample markdown content.", encoding="utf-8")

    doc = load_document(sample_file)
    assert doc.doc_id == "sample"
    assert doc.source == "sample.md"
    assert "Test Title" in doc.text
    assert doc.metadata["file_type"] == ".md"


def test_load_json_file(tmp_path: Path):
    sample_file = tmp_path / "data.json"
    data = {"company": "Enterprise AI", "employees": 500, "location": "Hanoi"}
    sample_file.write_text(json.dumps(data), encoding="utf-8")

    doc = load_document(sample_file)
    assert doc.doc_id == "data"
    assert "Enterprise AI" in doc.text
    assert doc.metadata["file_type"] == ".json"


def test_load_word_file(tmp_path: Path):
    import docx
    sample_file = tmp_path / "test.docx"
    doc_obj = docx.Document()
    doc_obj.add_heading("Chính sách bảo mật", level=1)
    doc_obj.add_paragraph("Nội dung chính sách bảo mật thông tin khách hàng.")
    doc_obj.save(sample_file)

    doc = load_document(sample_file)
    assert doc.doc_id == "test"
    assert "Chính sách bảo mật" in doc.text
    assert doc.metadata["file_type"] == ".docx"


def test_load_excel_csv_file(tmp_path: Path):
    sample_file = tmp_path / "report.csv"
    sample_file.write_text("Thang,DoanhThu\n1,1000\n2,1500\n", encoding="utf-8")

    doc = load_document(sample_file)
    assert doc.doc_id == "report"
    assert "DoanhThu" in doc.text
    assert doc.metadata["file_type"] == ".csv"


def test_load_image_synthetic(tmp_path: Path):
    """Test creating an image with text and running loader on it."""
    img_path = tmp_path / "test_chart.png"
    img = Image.new("RGB", (300, 100), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    d.text((10, 10), "Doanh thu Q1: 100 trieu", fill=(0, 0, 0))
    d.text((10, 30), "Doanh thu Q2: 250 trieu", fill=(0, 0, 0))
    img.save(img_path)

    doc = load_document(img_path)
    assert doc.doc_id == "test_chart"
    assert doc.source == "test_chart.png"
    assert doc.metadata["file_type"] == "image"
    assert len(doc.text) > 0
