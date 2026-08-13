"""
Thiết lập Bài toán Thực tế Doanh nghiệp (Concrete Business Use Case):
Hệ thống Trợ lý Tra cứu Pháp lý, Quy chế Nhân sự & Kế toán Tài chính Doanh nghiệp
(Enterprise HR, Finance & Compliance AI Assistant)

Bao gồm dữ liệu đa định dạng thực tế:
- 01_Quy_che_Thuong_va_Nhan_su_2026.docx (Word Document - Quy chế nhân sự)
- 02_Bang_Bieu_Phu_Cap_Cong_Tac_Phi.xlsx (Excel File - Bảng lương phụ cấp)
- 03_Quy_trinh_Thanh_Toan_Hoa_Don_VAT.md (Markdown - Quy trình tài chính)
- 04_Mau_Hoa_Don_Thanh_Toan_Scan.png (Synthetic Image - Mẫu hóa đơn scan)
"""
import os
import sys
import json
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from PIL import Image, ImageDraw, ImageFont
import docx
import openpyxl
import pandas as pd

sys.stdout.reconfigure(encoding="utf-8")

DOMAIN_DIR = Path("data/domain_enterprise_hr_finance")
DOMAIN_DIR.mkdir(parents=True, exist_ok=True)


def create_word_hr_policy():
    path = DOMAIN_DIR / "01_Quy_che_Thuong_va_Nhan_su_2026.docx"
    doc = docx.Document()
    doc.add_heading("QUY CHẾ NHÂN SỰ VÀ THƯỞNG HIỆU QUẢ CÔNG VIỆC 2026", level=1)

    doc.add_heading("Điều 1: Thời giờ làm việc và Nghỉ phép năm", level=2)
    doc.add_paragraph("1.1. Thời gian làm việc tiêu chuẩn là 8 giờ/ngày, từ thứ Hai đến thứ Sáu (8:00 - 17:00).")
    doc.add_paragraph("1.2. Mỗi nhân viên chính thức được hưởng 12 ngày phép năm có hưởng nguyên lương. Nhân viên có thâm niên trên 5 năm cứ mỗi 3 năm được cộng thêm 1 ngày phép.")

    doc.add_heading("Điều 2: Chính sách Thưởng và Phụ cấp Ngoài giờ (OT)", level=2)
    doc.add_paragraph("2.1. Phụ cấp làm thêm giờ (OT) ngày thường tính 150% lương giờ cơ bản. OT ngày nghỉ hàng tuần (Chủ nhật) tính 200%. OT ngày lễ tết tính 300%.")
    doc.add_paragraph("2.2. Nhân viên làm việc từ sau 22:00 được trợ cấp bữa ăn đêm 100.000 VNĐ/buổi.")
    doc.add_paragraph("2.3. Thưởng hiệu quả công việc (KPI) được xét duyệt theo quý, mức thưởng tối đa 3 tháng lương cho cá nhân đạt xếp loại Xuất sắc A+.")

    doc.add_heading("Điều 3: Quyền hạn và Trách nhiệm Bảo mật", level=2)
    doc.add_paragraph("Nhân viên vi phạm quy định bảo mật dữ liệu khách hàng sẽ bị xử lý kỷ luật sa thải và bồi thường thiệt hại tối thiểu 50.000.000 VNĐ.")

    doc.save(path)
    print(f"  [1/4] Created Word document: {path.name}")


def create_excel_allowance_table():
    path = DOMAIN_DIR / "02_Bang_Bieu_Phu_Cap_Cong_Tac_Phi.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Phụ cấp Công tác phí"

    headers = ["Chức danh / Cấp bậc", "Phụ cấp Tiền ăn/ngày (VNĐ)", "Khách sạn tối đa/đêm (VNĐ)", "Phụ cấp Xe đi lại/ngày (VNĐ)", "Hạn mức Vé máy bay"]
    ws.append(headers)

    rows = [
        ["Giám đốc / Ban Điều hành", 500000, 2500000, 400000, "Hạng Thương gia (Business)"],
        ["Trưởng phòng / Quản lý Cấp cao", 350000, 1500000, 250000, "Hạng Phổ thông linh hoạt"],
        ["Chuyên viên / Nhân viên chính thức", 250000, 900000, 150000, "Hạng Phổ thông (Economy)"],
        ["Thực tập sinh / Cộng tác viên", 150000, 500000, 100000, "Hạng Phổ thông tiết kiệm"],
    ]

    for r in rows:
        ws.append(r)

    wb.save(path)
    print(f"  [2/4] Created Excel spreadsheet: {path.name}")


def create_md_finance_policy():
    path = DOMAIN_DIR / "03_Quy_trinh_Thanh_Toan_Hoa_Don_VAT.md"
    content = """# QUY TRÌNH KẾ TOÁN VÀ THANH TOÁN HÓA ĐƠN VAT DOANH NGHIỆP

## 1. Quy định chung về Hóa đơn tài chính (VAT)
- Tất cả hóa đơn thanh toán có giá trị từ **5.000.000 VNĐ** trở lên bắt buộc phải là **Hóa đơn điện tử VAT hợp pháp** tra cứu được trên cổng Tổng cục Thuế.
- Hóa đơn mua hàng có giá trị từ **20.000.000 VNĐ** trở lên bắt buộc phải thực hiện thanh toán chuyển khoản từ tài khoản ngân hàng của Công ty (không thanh toán tiền mặt).

## 2. Quy trình Duyệt hồ sơ Thanh toán
1. **Bước 1:** Nhân viên đề xuất lập Tờ trình kèm Hợp đồng và Hóa đơn VAT gửi Kế toán viên kiểm tra trong vòng **2 ngày làm việc**.
2. **Bước 2:** Kế toán trưởng kiểm duyệt tính hợp lệ của chứng từ trong vòng **1 ngày làm việc**.
3. **Bước 3:** Giám đốc Tài chính (CFO) hoặc Tổng Giám đốc phê duyệt chi tiền.
4. **Bước 4:** Bộ phận Thủ quỹ / Ngân hàng thực hiện lệnh chuyển khoản vào thứ Ba và thứ Sáu hàng tuần.

## 3. Tạm ứng và Hoàn ứng Công tác phí
- Hạn mức tạm ứng tối đa cho mỗi chuyến công tác là **30.000.000 VNĐ**.
- Trong vòng **5 ngày làm việc** sau khi kết thúc chuyến công tác, nhân viên phải nộp đầy đủ vé máy bay, hóa đơn khách sạn và bảng kê chi tiết để hoàn ứng.
"""
    path.write_text(content, encoding="utf-8")
    print(f"  [3/4] Created Markdown document: {path.name}")


def create_scanned_receipt_image():
    path = DOMAIN_DIR / "04_Mau_Hoa_Don_Thanh_Toan_Scan.png"
    img = Image.new("RGB", (650, 250), color=(255, 255, 255))
    d = ImageDraw.Draw(img)

    d.rectangle([10, 10, 640, 240], outline=(30, 64, 175), width=3)
    d.text((25, 25), "PHIẾU ĐỀ XUẤT XÁC NHẬN CHI PHÍ TIẾP KHÁCH SCAN", fill=(30, 64, 175))
    d.text((25, 60), "Mã hồ sơ: REF-2026-9982", fill=(15, 23, 42))
    d.text((25, 90), "Nội dung: Chi phí tiếp đối tác dự án AI Enterprise", fill=(15, 23, 42))
    d.text((25, 120), "Số tiền chi trả: 8.500.000 VNĐ (Tám triệu năm trăm nghìn đồng)", fill=(15, 23, 42))
    d.text((25, 150), "Trạng thái: Đã có hóa đơn VAT điện tử hợp lệ", fill=(22, 163, 74))
    d.text((25, 180), "Người phê duyệt chi: Giám đốc Tài chính (Đã ký điện tử)", fill=(15, 23, 42))

    img.save(path)
    print(f"  [4/4] Created Scanned Image Document: {path.name}")


def ingest_concrete_domain():
    print("\n========================================================")
    print("📥 NẠP VÀ INDEX BÀI TOÁN THỰC TẾ: HR & FINANCE ASSISTANT")
    print("========================================================")
    from app.ingestion.pipeline import IngestionPipeline
    from app.graph.agentic_rag import get_hybrid_retriever

    pipeline = IngestionPipeline()
    res = pipeline.ingest_directory(DOMAIN_DIR, reset=True)
    get_hybrid_retriever().reload_index()
    print(f"  -> Successfully ingested {res.get('documents_loaded')} documents into {res.get('chunks_generated')} chunks.")


if __name__ == "__main__":
    create_word_hr_policy()
    create_excel_allowance_table()
    create_md_finance_policy()
    create_scanned_receipt_image()
    ingest_concrete_domain()
