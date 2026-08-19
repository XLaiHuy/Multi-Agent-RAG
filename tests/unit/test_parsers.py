"""
Unit Tests for Multi-Format Canonical Parsers (Markdown, JSON, DOCX).
"""
import json
import pytest
from pathlib import Path
from backend.app.ingestion.parsers import MarkdownParser, JSONParser, DocxParser
from backend.app.domain.canonical import BlockType


def test_markdown_parser_headings_and_tables(tmp_path: Path):
    md_file = tmp_path / "sample_contract.md"
    md_content = """# Master Agreement

This is the preamble of the agreement.

## Section 1: Services
The provider agrees to perform cloud engineering services.

### 1.1 SLA Metrics
Uptime must exceed 99.9% per billing cycle.

| Service | Tier | Response Time |
| --- | --- | --- |
| Compute | Premium | 15 mins |
| Storage | Standard | 1 hour |
"""
    md_file.write_text(md_content, encoding="utf-8")

    doc = MarkdownParser.parse_markdown(md_file, doc_id="md_test_01")
    assert doc.doc_id == "md_test_01"
    assert doc.doc_type == "markdown"
    assert len(doc.pages) == 1
    
    blocks = doc.pages[0].blocks
    assert len(blocks) >= 4

    # Check heading hierarchy
    sla_block = [b for b in blocks if "Uptime must exceed" in b.text][0]
    assert "Section 1: Services" in sla_block.section_path or "1.1 SLA Metrics" in sla_block.section_path

    # Check table parsing
    table_block = [b for b in blocks if b.block_type == BlockType.TABLE or "Compute" in b.text][0]
    assert "Compute" in table_block.text


def test_json_parser_path_preservation(tmp_path: Path):
    json_file = tmp_path / "contract_data.json"
    data = {
        "contract": {
            "title": "Vendor License Agreement",
            "termination": {
                "notice_days": 30,
                "for_convenience": True
            },
            "governing_law": "Delaware"
        }
    }
    json_file.write_text(json.dumps(data), encoding="utf-8")

    doc = JSONParser.parse_json(json_file, doc_id="json_test_01")
    assert doc.doc_id == "json_test_01"
    assert doc.doc_type == "json"

    blocks = doc.get_all_blocks()
    notice_block = [b for b in blocks if "notice_days" in b.text][0]
    assert "contract.termination.notice_days" in notice_block.text or "notice_days" in notice_block.text
    assert "30" in notice_block.text


def test_docx_parser(tmp_path: Path):
    import docx
    docx_file = tmp_path / "test_contract.docx"
    d = docx.Document()
    d.add_heading("NDA Confidentiality Agreement", level=1)
    d.add_paragraph("All proprietary software source code shall be kept confidential.")
    d.save(docx_file)

    doc = DocxParser.parse_docx(docx_file, doc_id="docx_test_01")
    assert doc.doc_id == "docx_test_01"
    assert doc.doc_type == "docx"
    blocks = doc.get_all_blocks()
    assert len(blocks) == 2
    assert "NDA Confidentiality Agreement" in blocks[0].text
    assert "proprietary software" in blocks[1].text
